from io import BytesIO
from unittest.mock import Mock
from zipfile import ZIP_STORED, ZipFile

import pytest
from docx import Document
from fastapi.testclient import TestClient
from starlette.datastructures import UploadFile

from app import api_service, uploads
from app.api import create_app, get_workflow_runner
from app.api_service import ProviderConfigurationError, ProviderServiceError
from app.state import build_initial_state


def docx_bytes(text="Synthetic resume text", padding=0):
    buffer = BytesIO()
    document = Document()
    if text:
        document.add_paragraph(text)
    document.save(buffer)
    if padding:
        with ZipFile(buffer, "a", compression=ZIP_STORED) as archive:
            archive.writestr("padding.bin", b"x" * padding)
    return buffer.getvalue()


def completed_state():
    state = build_initial_state("AI jobs", "PRIVATE_RESUME_SENTINEL")
    state.update(role="AI Engineer", location="Remote", employment_type="Full-time")
    state["candidate_profile"] = {
        "summary": "Experienced Python engineer", "years_experience": 10,
        "resume_text": "PRIVATE_RESUME_SENTINEL", "secret": "SECRET_SENTINEL",
    }
    verified = {
        "title": "Engineer", "company": "Example", "location": "Remote",
        "match_score": 90, "preliminary_match_score": 80,
        "verification_status": "verified", "analysis_type": "verified",
        "strengths": ["Python"], "missing_skills": ["Go"],
        "confidence": "High", "recommendation": "Strong Apply",
        "source_url": "https://example.com/job/1",
        "verified_url": "https://example.com/careers/1",
        "description_url": "https://example.com/careers/1",
        "description": "INTERNAL_DESCRIPTION_SENTINEL",
    }
    preliminary = {
        **verified, "title": "Backend Engineer", "match_score": 78,
        "verification_status": "service_error", "analysis_type": "preliminary",
    }
    state.update(
        jobs=[verified, preliminary], analyses=[verified, preliminary],
        verification_candidates=[verified, preliminary],
        verified_jobs=[verified, preliminary], selected_jobs=[verified],
        final_ranked_jobs=[verified, preliminary],
    )
    state["secret"] = "SECRET_SENTINEL"
    return state


@pytest.fixture(autouse=True)
def enable_live_search(monkeypatch):
    monkeypatch.setenv("ENABLE_LIVE_SEARCH", "true")


@pytest.fixture
def application(monkeypatch):
    monkeypatch.delenv("CORS_ORIGINS", raising=False)
    return create_app()


@pytest.fixture
def runner(application):
    mock = Mock(return_value=completed_state())
    application.dependency_overrides[get_workflow_runner] = lambda: mock
    return mock


def post(application, content=None, filename="resume.docx", media_type=uploads.DOCX_MEDIA_TYPE, data=None):
    with TestClient(application, raise_server_exceptions=False) as client:
        return client.post(
            "/api/v1/job-search",
            files={"resume": (filename, docx_bytes() if content is None else content, media_type)},
            data={"search_request": "  Find remote AI jobs  "} if data is None else data,
        )


def test_health_without_credentials_or_provider_initialization(application, monkeypatch):
    from app import nodes
    from app.tools import web_search, job_search

    forbidden = Mock(side_effect=AssertionError("Health initialized a provider"))
    for name in ("OPENAI_API_KEY", "JOOBLE_API_KEY", "TAVILY_API_KEY"):
        monkeypatch.delenv(name, raising=False)
    monkeypatch.setattr(nodes, "ChatOpenAI", forbidden)
    monkeypatch.setattr(web_search, "TavilyClient", forbidden)
    monkeypatch.setattr(job_search.requests, "post", forbidden)
    with TestClient(application) as client:
        response = client.get("/health")
    assert response.status_code == 200
    assert response.json() == {"status": "ok"}
    forbidden.assert_not_called()


def test_search_returns_public_contract_only(application, runner):
    response = post(application)
    assert response.status_code == 200, response.text
    runner.assert_called_once_with("Find remote AI jobs", "Synthetic resume text")
    body = response.json()
    assert set(body) == {"criteria", "candidate_profile", "ranked_jobs", "run_summary"}
    assert body["criteria"] == {
        "role": "AI Engineer", "location": "Remote", "employment_type": "Full-time", "days_old": 7,
    }
    assert body["candidate_profile"] == {"summary": "Experienced Python engineer", "years_experience": 10}
    verified, preliminary = body["ranked_jobs"]
    assert verified["verified_match_score"] == 90
    assert verified["preliminary_match_score"] == 80
    assert preliminary["verified_match_score"] is None
    assert preliminary["preliminary_match_score"] == 78
    assert preliminary["source_urls"]["verified"] is None
    assert verified["source_urls"]["original"] == "https://example.com/job/1"
    assert verified["strengths"] == ["Python"]
    assert verified["missing_skills"] == ["Go"]
    assert body["run_summary"] == {
        "status": "partial", "jobs_found": 2, "jobs_analyzed": 2,
        "verification_attempted": 2, "verified_jobs": 1, "preliminary_jobs": 1,
        "selected_jobs": 1, "returned_jobs": 2,
        "warnings": ["Some jobs could not be verified; preliminary scores were retained."],
    }
    for sentinel in ("PRIVATE_RESUME_SENTINEL", "SECRET_SENTINEL", "INTERNAL_DESCRIPTION_SENTINEL"):
        assert sentinel not in response.text


def test_api_extracts_paragraphs_and_tables_before_running_graph(application, runner):
    document = Document()
    document.add_paragraph("Overview")
    document.add_table(rows=1, cols=2).cell(0, 0).text = "Python experience"
    document.add_paragraph("Education")
    buffer = BytesIO()
    document.save(buffer)
    response = post(application, content=buffer.getvalue())
    assert response.status_code == 200
    runner.assert_called_once_with("Find remote AI jobs", "Overview\nPython experience\nEducation")


def test_verified_source_does_not_promote_explicit_preliminary_analysis(application, runner):
    state = completed_state()
    state["final_ranked_jobs"][0]["analysis_type"] = "preliminary"
    runner.return_value = state
    result = post(application).json()["ranked_jobs"][0]
    assert result["verified_match_score"] is None
    assert result["analysis_type"] == "preliminary"


@pytest.mark.parametrize("status", [None, "", "unknown", "not_needed", "VERIFIED", True, "missing"])
def test_api_never_promotes_unknown_status(application, runner, status):
    state = completed_state()
    state["final_ranked_jobs"][0]["verification_status"] = status
    if status == "missing":
        del state["final_ranked_jobs"][0]["verification_status"]
    runner.return_value = state
    result = post(application).json()["ranked_jobs"][0]
    assert result["analysis_type"] == "preliminary"
    assert result["verified_match_score"] is None
    assert result["preliminary_match_score"] == 90


def test_empty_results_return_structured_success(application, runner):
    state = completed_state()
    for key in ("jobs", "analyses", "verification_candidates", "verified_jobs", "selected_jobs", "final_ranked_jobs"):
        state[key] = []
    runner.return_value = state
    response = post(application)
    assert response.status_code == 200
    assert response.json()["ranked_jobs"] == []
    assert response.json()["run_summary"]["status"] == "completed"


@pytest.mark.parametrize("data", [
    {}, {"search_request": " "}, {"search_request": "x" * 2001},
    {"search_request": "AI jobs", "verification_limit": "2"},
])
def test_invalid_form_is_safe(application, runner, data):
    response = post(application, data=data)
    assert response.status_code == 422
    assert response.json()["error"]["code"] == "invalid_request"
    runner.assert_not_called()


def test_missing_upload(application, runner):
    with TestClient(application) as client:
        response = client.post("/api/v1/job-search", data={"search_request": "AI jobs"})
    assert response.status_code == 422
    runner.assert_not_called()


@pytest.mark.parametrize("filename, media_type", [
    ("resume.pdf", "application/pdf"),
    ("resume.docx", "text/plain"),
    ("resume.docm", uploads.DOCX_MEDIA_TYPE),
])
def test_wrong_file_type(application, runner, filename, media_type):
    response = post(application, filename=filename, media_type=media_type)
    assert response.status_code == 415
    runner.assert_not_called()


@pytest.mark.parametrize("content", [b"", b"not a docx", docx_bytes("")], ids=["empty", "garbage", "blank-docx"])
def test_unreadable_docx(application, runner, content):
    response = post(application, content=content)
    assert response.status_code == 422
    assert response.json()["error"]["code"] == "unreadable_docx"
    runner.assert_not_called()


def test_expanded_archive_limit(application, runner, monkeypatch):
    monkeypatch.setattr(uploads, "MAX_EXPANDED_BYTES", 1)
    response = post(application)
    assert response.status_code == 422
    runner.assert_not_called()


@pytest.mark.parametrize("size", [uploads.MAX_UPLOAD_BYTES + 1, uploads.MAX_REQUEST_BYTES + 1])
def test_oversized_upload(application, runner, size):
    response = post(application, content=b"x" * size)
    assert response.status_code == 413
    runner.assert_not_called()


@pytest.mark.parametrize("error, status, code", [
    (ProviderConfigurationError("SECRET_SENTINEL"), 503, "provider_not_configured"),
    (ProviderServiceError("SECRET_SENTINEL"), 502, "provider_failure"),
    (RuntimeError("SECRET_SENTINEL"), 500, "internal_error"),
])
def test_errors_do_not_leak_details(application, runner, error, status, code):
    runner.side_effect = error
    response = post(application)
    assert response.status_code == status
    assert response.json()["error"]["code"] == code
    assert "SECRET_SENTINEL" not in response.text
    assert "Traceback" not in response.text


@pytest.mark.parametrize("mode", ["success", "failure", "unreadable", "invalid_request", "oversized", "wrong_type"])
def test_spooled_upload_is_closed_on_every_path(application, runner, monkeypatch, mode):
    closed = []
    original_close = UploadFile.close

    async def track_close(self):
        await original_close(self)
        closed.append(self.file)

    monkeypatch.setattr(UploadFile, "close", track_close)
    content = docx_bytes(padding=1100 * 1024)
    if mode == "failure":
        runner.side_effect = RuntimeError("fail")
    if mode == "unreadable":
        content = b"bad-docx" * (150 * 1024)
    if mode == "oversized":
        content = b"x" * (uploads.MAX_UPLOAD_BYTES + 1)
    response = post(
        application, content=content, data={} if mode == "invalid_request" else None,
        filename="bad.pdf" if mode == "wrong_type" else "resume.docx",
    )
    assert response.status_code == {
        "success": 200, "failure": 500, "unreadable": 422,
        "invalid_request": 422, "oversized": 413, "wrong_type": 415,
    }[mode]
    assert closed
    assert all(file.closed for file in closed)
    assert any(getattr(file, "_rolled", False) for file in closed)


def test_duplicate_uploads_are_all_closed(application, runner, monkeypatch):
    closed = []
    original_close = UploadFile.close

    async def track_close(self):
        await original_close(self)
        closed.append(self.file)

    monkeypatch.setattr(UploadFile, "close", track_close)
    content = docx_bytes(padding=1100 * 1024)
    with TestClient(application) as client:
        response = client.post(
            "/api/v1/job-search", data={"search_request": "AI jobs"},
            files=[("resume", (name, content, uploads.DOCX_MEDIA_TYPE))
                   for name in ("first.docx", "second.docx")],
        )
    assert response.status_code == 422
    assert len({id(file) for file in closed}) == 2
    assert all(file.closed for file in closed)
    runner.assert_not_called()


def test_multipart_parser_failure_closes_created_uploads(application, runner, monkeypatch):
    from starlette import formparsers
    created = []
    original_spooled_file = formparsers.SpooledTemporaryFile

    def track_file(*args, **kwargs):
        file = original_spooled_file(*args, **kwargs)
        created.append(file)
        return file

    monkeypatch.setattr(formparsers, "SpooledTemporaryFile", track_file)
    body = (
        b'--boundary\r\nContent-Disposition: form-data; name="resume"; filename="resume.docx"\r\n'
        b'Content-Type: application/octet-stream\r\n\r\npartial file data\r\n'
        b'--boundary\r\nINVALID HEADER WITHOUT COLON\r\n\r\nSECRET_SENTINEL\r\n--boundary--\r\n'
    )
    with TestClient(application) as client:
        response = client.post("/api/v1/job-search", content=body,
                               headers={"Content-Type": "multipart/form-data; boundary=boundary"})
    assert response.status_code == 400
    assert created and all(file.closed for file in created)
    assert "SECRET_SENTINEL" not in response.text
    runner.assert_not_called()


def test_required_configuration_checked_before_graph(application, monkeypatch):
    from app.graph import graph
    invoke = Mock(side_effect=AssertionError("Graph should not run"))
    monkeypatch.setattr(graph, "invoke", invoke)
    for name in ("OPENAI_API_KEY", "JOOBLE_API_KEY", "TAVILY_API_KEY"):
        monkeypatch.delenv(name, raising=False)
    response = post(application)
    assert response.status_code == 503
    invoke.assert_not_called()


def test_service_invokes_existing_graph(monkeypatch):
    from app.graph import graph
    for name in ("OPENAI_API_KEY", "JOOBLE_API_KEY", "TAVILY_API_KEY"):
        monkeypatch.setenv(name, "fake-test-value")
    invoke = Mock(return_value=completed_state())
    monkeypatch.setattr(graph, "invoke", invoke)
    result = api_service.run_workflow("AI jobs", "Synthetic resume text")
    assert result["role"] == "AI Engineer"
    assert invoke.call_args.args[0] == build_initial_state("AI jobs", "Synthetic resume text")


def test_provider_failure_classification(application, monkeypatch):
    import requests
    from app.graph import graph
    for name in ("OPENAI_API_KEY", "JOOBLE_API_KEY", "TAVILY_API_KEY"):
        monkeypatch.setenv(name, "fake-test-value")
    monkeypatch.setattr(graph, "invoke", Mock(side_effect=requests.Timeout("SECRET_SENTINEL")))
    response = post(application)
    assert response.status_code == 502
    assert "SECRET_SENTINEL" not in response.text


def test_cors_uses_only_configured_origins(monkeypatch):
    monkeypatch.setenv("CORS_ORIGINS", "http://localhost:5173")
    with TestClient(create_app()) as client:
        allowed = client.options("/api/v1/job-search", headers={
            "Origin": "http://localhost:5173", "Access-Control-Request-Method": "POST",
        })
        rejected = client.options("/api/v1/job-search", headers={
            "Origin": "https://untrusted.example", "Access-Control-Request-Method": "POST",
        })
    assert allowed.headers["access-control-allow-origin"] == "http://localhost:5173"
    assert "access-control-allow-credentials" not in allowed.headers
    assert "access-control-allow-origin" not in rejected.headers


def test_wildcard_cors_rejected(monkeypatch):
    monkeypatch.setenv("CORS_ORIGINS", "*")
    with pytest.raises(ValueError):
        create_app()


def test_cors_headers_present_on_internal_failure(monkeypatch):
    monkeypatch.setenv("CORS_ORIGINS", "http://localhost:5173")
    application = create_app()
    runner = Mock(side_effect=RuntimeError("SECRET_SENTINEL"))
    application.dependency_overrides[get_workflow_runner] = lambda: runner
    with TestClient(application) as client:
        response = client.post(
            "/api/v1/job-search",
            headers={"Origin": "http://localhost:5173"},
            files={"resume": ("resume.docx", docx_bytes(), uploads.DOCX_MEDIA_TYPE)},
            data={"search_request": "AI jobs"},
        )
    assert response.status_code == 500
    assert response.headers["access-control-allow-origin"] == "http://localhost:5173"
    assert "SECRET_SENTINEL" not in response.text


def test_openapi_documents_contract(application):
    schema = application.openapi()
    assert set(schema["paths"]) == {"/health", "/api/v1/demo", "/api/v1/job-search"}
    assert "multipart/form-data" in schema["paths"]["/api/v1/job-search"]["post"]["requestBody"]["content"]
    assert "JobSearchResponse" in schema["components"]["schemas"]


@pytest.mark.parametrize("url", [
    "https://user:SECRET_SENTINEL@example.com/job",
    "https://example.com/job?api_key=SECRET_SENTINEL",
    "https://example.com/job#access_token=SECRET_SENTINEL",
    "https://example.com/job?X-Amz-Signature=SECRET_SENTINEL",
    "https://example.com/job?API-KEY=SECRET_SENTINEL",
    "javascript:SECRET_SENTINEL",
])
def test_source_urls_cannot_expose_credentials(application, runner, url):
    state = completed_state()
    for field in ("source_url", "verified_url", "description_url"):
        state["final_ranked_jobs"][0][field] = url
    runner.return_value = state
    response = post(application)
    assert response.status_code == 200
    assert response.json()["ranked_jobs"][0]["source_urls"] == {
        "original": None, "verified": None, "description": None,
    }
    assert "SECRET_SENTINEL" not in response.text


@pytest.mark.parametrize("path", ["/api/v1/job-search", "/api/v1/job-search/", "/prefix/api/v1/job-search"])
def test_chunked_body_limit_applies_without_content_length(monkeypatch, path):
    import asyncio
    from app import api

    monkeypatch.setattr(api, "MAX_REQUEST_BYTES", 4)
    messages = iter([
        {"type": "http.request", "body": b"abc", "more_body": True},
        {"type": "http.request", "body": b"def", "more_body": False},
    ])
    sent = []

    async def receive():
        return next(messages)

    async def send(message):
        sent.append(message)

    async def forbidden_app(scope, receive, send):
        pytest.fail("Oversized multipart body reached the parser")

    asyncio.run(api.BodyLimitMiddleware(forbidden_app, live_enabled=True)(
        {"type": "http", "method": "POST", "path": path, "headers": []},
        receive, send,
    ))
    assert sent[0]["status"] == 413


def test_zero_verification_limit_does_not_require_tavily(monkeypatch):
    from app.graph import graph
    monkeypatch.setenv("MAX_VERIFICATION_JOBS", "0")
    monkeypatch.setenv("OPENAI_API_KEY", "fake-test-value")
    monkeypatch.setenv("JOOBLE_API_KEY", "fake-test-value")
    monkeypatch.delenv("TAVILY_API_KEY", raising=False)
    monkeypatch.setattr(graph, "invoke", Mock(return_value={}))
    assert api_service.run_workflow("AI jobs", "Synthetic resume text") == {}


def test_malformed_multipart_does_not_echo_parser_details(application, runner):
    with TestClient(application) as client:
        response = client.post(
            "/api/v1/job-search", content=b"SECRET_SENTINEL",
            headers={"Content-Type": "multipart/form-data"},
        )
    assert response.status_code == 400
    assert "SECRET_SENTINEL" not in response.text
    runner.assert_not_called()
