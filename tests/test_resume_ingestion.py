from io import BytesIO
import struct
from unittest.mock import Mock
from zipfile import ZIP_DEFLATED, ZIP_STORED, ZipExtFile, ZipFile
from zlib import crc32

from docx import Document
import pytest

from app import uploads
from app.tools.resume_reader import read_docx_resume


def save_document(document):
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def mixed_document():
    document = Document()
    document.add_paragraph("  Candidate overview  ")
    document.add_paragraph("")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Experience"
    table.cell(0, 1).text = "Senior Engineer"
    table.cell(1, 0).text = "Skills"
    table.cell(1, 1).text = "Python"
    document.add_paragraph("Project highlights")
    document.add_table(rows=1, cols=1).cell(0, 0).text = "Built search systems"
    document.add_paragraph("Education")
    document.sections[0].header.paragraphs[0].text = "HEADER_ONLY"
    document.sections[0].footer.paragraphs[0].text = "FOOTER_ONLY"
    return document


def test_mixed_paragraphs_and_tables_preserve_body_order():
    text = read_docx_resume(BytesIO(save_document(mixed_document())))
    assert text.splitlines() == [
        "Candidate overview", "Experience", "Senior Engineer", "Skills",
        "Python", "Project highlights", "Built search systems", "Education",
    ]


def test_merged_cells_are_not_repeated_but_distinct_equal_text_is_retained():
    document = Document()
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "Merged heading"
    table.cell(1, 0).merge(table.cell(2, 0)).text = "Merged experience"
    table.cell(1, 1).text = "Python"
    table.cell(2, 1).text = "Python"
    assert read_docx_resume(BytesIO(save_document(document))).splitlines() == [
        "Merged heading", "Merged experience", "Python", "Python",
    ]


def test_nested_table_stays_between_cell_paragraphs():
    document = Document()
    document.add_paragraph("Before table")
    cell = document.add_table(rows=1, cols=1).cell(0, 0)
    cell.text = "Before nested table"
    cell.add_table(rows=1, cols=1).cell(0, 0).text = "Nested skills"
    cell.add_paragraph("After nested table")
    document.add_paragraph("After table")
    assert read_docx_resume(BytesIO(save_document(document))).splitlines() == [
        "Before table", "Before nested table", "Nested skills",
        "After nested table", "After table",
    ]


def test_file_path_input_is_still_supported(tmp_path):
    path = tmp_path / "synthetic-resume.docx"
    mixed_document().save(path)
    assert "Senior Engineer" in read_docx_resume(str(path))


def test_table_only_resume_is_readable():
    document = Document()
    document.add_table(rows=1, cols=1).cell(0, 0).text = "Python engineer"
    assert uploads.extract_uploaded_resume(save_document(document)) == "Python engineer"


def test_header_only_resume_is_rejected():
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Header, not body text"
    with pytest.raises(uploads.UnreadableResume):
        uploads.extract_uploaded_resume(save_document(document))


def test_duplicate_archive_members_are_rejected():
    buffer = BytesIO(save_document(mixed_document()))
    with ZipFile(buffer, "a") as archive, pytest.warns(UserWarning, match="Duplicate name"):
        archive.writestr("word/document.xml", "duplicate")
    with pytest.raises(uploads.UnreadableResume):
        uploads.extract_uploaded_resume(buffer.getvalue())


def test_archive_entry_count_limit(monkeypatch):
    monkeypatch.setattr(uploads, "MAX_ARCHIVE_FILES", 1)
    with pytest.raises(uploads.UnreadableResume):
        uploads.extract_uploaded_resume(save_document(mixed_document()))


def test_compressed_zip_with_forged_size_is_repacked_before_docx_reads(monkeypatch):
    buffer = BytesIO(save_document(mixed_document()))
    with ZipFile(buffer, "a", compression=ZIP_DEFLATED) as archive:
        archive.writestr("padding.bin", b"x" * (2 * 1024 * 1024))
    content = bytearray(buffer.getvalue())
    # Forge a central-directory size and CRC for a short prefix of a large
    # compressed member. Never pass this original ZIP to unbounded reader calls.
    central_offset = content.rfind(b"padding.bin") - 46
    assert content[central_offset:central_offset + 4] == b"PK\x01\x02"
    struct.pack_into("<I", content, central_offset + 24, 16)
    struct.pack_into("<I", content, central_offset + 16, crc32(b"x" * 16))
    original_read = ZipExtFile.read
    original_reader = uploads.read_docx_resume
    compressed_reads = []

    def bounded_read(self, n=-1):
        if self._compress_type != ZIP_STORED:
            assert 0 < n <= 64 * 1024
            compressed_reads.append(n)
        return original_read(self, n)

    def inspect_normalized(stream):
        with ZipFile(stream) as archive:
            assert all(entry.compress_type == ZIP_STORED for entry in archive.infolist())
            assert archive.read("padding.bin") == b"x" * 16
        stream.seek(0)
        return original_reader(stream)

    monkeypatch.setattr(ZipExtFile, "read", bounded_read)
    monkeypatch.setattr(uploads, "read_docx_resume", inspect_normalized)
    assert "Senior Engineer" in uploads.extract_uploaded_resume(bytes(content))
    assert compressed_reads


def test_cli_still_loads_resume_runs_graph_and_prints_report(monkeypatch, capsys):
    import main
    from app.graph import graph
    from app.tools import resume_reader

    load_env = Mock()
    reader = Mock(return_value="Synthetic resume")
    invoke = Mock(return_value={"final_report": "Synthetic report"})
    monkeypatch.setattr(main, "load_dotenv", load_env)
    monkeypatch.setattr(resume_reader, "read_docx_resume", reader)
    monkeypatch.setattr(graph, "invoke", invoke)
    main.main()
    load_env.assert_called_once()
    reader.assert_called_once_with("data/resume.docx")
    assert invoke.call_args.args[0]["resume_text"] == "Synthetic resume"
    assert invoke.call_args.args[0]["search_request"] == (
        "Find remote Senior AI Engineer jobs from the last 7 days"
    )
    assert capsys.readouterr().out == "Synthetic report\n"
